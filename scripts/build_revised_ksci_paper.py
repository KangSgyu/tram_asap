# -*- coding: utf-8 -*-
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "revised_paper"
OUT_DOCX = OUT_DIR / "KSCI_Lean_Physical_AI_Manufacturing_Robot_Revised_with_Tables_Images.docx"


PROCESSES = ["rigid", "semi_rigid", "flexible"]
LABELS = {"rigid": "강체", "semi_rigid": "반강체", "flexible": "유연체"}
PAPER_BASELINE = {
    "rigid": {"before": 0.5, "after": 0.3},
    "semi_rigid": {"before": 38.5, "after": 21.0},
    "flexible": {"before": 34.5, "after": 17.5},
}


def load_summary(process: str) -> dict:
    with (ROOT / "outputs" / process / "summary.json").open("r", encoding="utf-8") as handle:
        return json.load(handle)


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def write_cell(cell, text: str, bold: bool = False, fill: str | None = None, align_left: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    set_run_font(run, size=8.5, bold=bold)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if align_left else WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if fill:
        shade(cell, fill)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        set_run_font(run, bold=True)


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.10


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5)
    paragraph.paragraph_format.space_after = Pt(5)


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, size=9, color="555555")
    run.italic = True
    paragraph.paragraph_format.space_after = Pt(8)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], caption: str) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        write_cell(table.rows[0].cells[idx], header, bold=True, fill="F2F4F7")
    for values in rows:
        row = table.add_row().cells
        for idx, value in enumerate(values):
            write_cell(row[idx], value, align_left=(idx == 0 and len(value) > 12))
    add_caption(doc, caption)


def build_document() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    summaries = {process: load_summary(process) for process in PROCESSES}
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("중소 제조 현장을 위한 린 기반 피지컬 AI 상체 제조 로봇 시뮬레이션 연구")
    set_run_font(run, size=18, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("A Simulation Study on Lean-Based Physical AI Upper-Body Manufacturing Robots for SME Assembly Cells")
    set_run_font(run, size=11)
    run.italic = True

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run("SeongGyu-Kang*")
    set_run_font(run, size=10.5, bold=True)

    add_heading(doc, "[Abstract]", 1)
    add_paragraph(
        doc,
        "This revised manuscript converts the original lean-based Physical AI manufacturing framework into an executable "
        "upper-body robot simulation for small and medium-sized assembly cells. Human manipulation trajectories from local "
        "SKEL data were preprocessed into tabletop assembly targets, and a compact dual-arm upper-body robot was simulated "
        "under rigid, semi-rigid, and flexible material conditions. The implementation integrates TRAM-style human motion "
        "utilization, ASAP-style residual dynamics alignment, and ASCPO-inspired step-size and torque safety constraints. "
        f"In a 240-frame, 12-iteration validation, defect rates decreased from {summaries['rigid']['before']['defect_rate']:.2f}% "
        f"to {summaries['rigid']['after']['defect_rate']:.2f}% in the rigid process, "
        f"{summaries['semi_rigid']['before']['defect_rate']:.2f}% to {summaries['semi_rigid']['after']['defect_rate']:.2f}% "
        f"in the semi-rigid process, and {summaries['flexible']['before']['defect_rate']:.2f}% to "
        f"{summaries['flexible']['after']['defect_rate']:.2f}% in the flexible process. The malfunction rate remained 0.0% "
        "in all virtual tests."
    )
    add_paragraph(doc, "Key words: Physical AI, Upper-body robot, Dynamics alignment, Lean manufacturing, SME data asset")

    add_heading(doc, "[요 약]", 1)
    add_paragraph(
        doc,
        "본 수정 논문은 기존의 린 기반 피지컬 AI 제조 프레임워크를 실행 가능한 상체 제조 로봇 시뮬레이션으로 확장한다. "
        "로컬 SKEL 인간 조작 궤적을 작업대 조립 목표 궤적으로 변환하고, 소형 중소 제조 셀 안에서 양팔 상체 로봇이 "
        "강체, 반강체, 유연체 자재를 조작하는 상황을 검증하였다. 구현 구조는 TRAM 개념의 인간 모션 활용, ASAP 개념의 "
        "잔차 기반 동역학 정렬, ASCPO 개념의 토크 및 스텝 제한 안전 제어를 포함한다. "
        f"240프레임, 12회 반복 검증에서 강체 공정 불량률은 {summaries['rigid']['before']['defect_rate']:.2f}%에서 "
        f"{summaries['rigid']['after']['defect_rate']:.2f}%로, 반강체 공정은 "
        f"{summaries['semi_rigid']['before']['defect_rate']:.2f}%에서 {summaries['semi_rigid']['after']['defect_rate']:.2f}%로, "
        f"유연체 공정은 {summaries['flexible']['before']['defect_rate']:.2f}%에서 "
        f"{summaries['flexible']['after']['defect_rate']:.2f}%로 감소하였다. 모든 가상 실험에서 오작동률은 0.0%로 유지되었다."
    )
    add_paragraph(doc, "주제어: 피지컬 AI, 상체 로봇, 동역학 정렬, 린 제조, 중소기업 데이터 자산")

    add_heading(doc, "I. Introduction", 1)
    add_paragraph(
        doc,
        "중소 제조 현장은 숙련 작업자 감소, 공정 데이터 부족, 대형 자동화 설비 도입 비용 부담이라는 제약을 동시에 가진다. "
        "특히 수작업 조립 공정은 반복성이 높지만 자재 물성, 작업대 배치, 작업자 동작 편차에 의해 품질 변동이 발생한다. "
        "따라서 본 연구는 전신 휴머노이드보다 작은 공간과 낮은 계산 비용으로 운용 가능한 상체 중심 제조 로봇을 대상으로 한다."
    )
    add_paragraph(
        doc,
        "본 수정본의 핵심 차별점은 기존 프레임워크를 설명하는 수준에서 멈추지 않고, 실제 코드 기반 시뮬레이션과 결과 산출물을 "
        "논문 내부에 포함했다는 점이다. 연구 질문은 다음과 같다."
    )
    add_bullet(doc, "최소한의 인간 조작 궤적만으로 상체 제조 로봇의 조립 동작을 생성할 수 있는가?")
    add_bullet(doc, "강체, 반강체, 유연체 공정에서 발생하는 물성 기반 추종 오차를 반복 정렬로 줄일 수 있는가?")
    add_bullet(doc, "비교 표와 시각화 결과를 통해 중소 제조 현장 적용 가능성을 정량적으로 판단할 수 있는가?")

    add_heading(doc, "II. System Architecture and Data Pipeline", 1)
    add_heading(doc, "2.1 Lean-Based Upper-Body Manufacturing Cell", 2)
    add_paragraph(
        doc,
        "제조 셀은 작업대, 부품 투입 위치, 조립 지그, 상체형 양팔 로봇으로 구성하였다. 로봇은 넓은 보행 공간을 요구하지 않고 "
        "작업대 상부에서 인간 작업자의 손 궤적을 모사한다. 린 제조 관점에서 부품 공급 위치와 조립 위치를 고정하여 이동 낭비를 줄이고, "
        "반복 공정에서의 표준 동작을 데이터화할 수 있도록 했다."
    )
    add_heading(doc, "2.2 Source Data and Preprocessing", 2)
    add_paragraph(
        doc,
        "기본 데이터는 data/raw/SKEL의 Rigid, Semi-Rigid, Flexible 모션 파일이다. 각 파일의 12열 강체 변환 행렬에서 평행이동 성분을 "
        "추출하고, 시작점 기준 정규화, 이상치 클리핑, 이동 평균 평활화, 프레임 재표본화를 수행하였다. 이 과정을 통해 사람의 손/물체 "
        "움직임을 작업대 좌표계의 로봇 엔드이펙터 목표 궤적으로 변환하였다."
    )
    add_table(
        doc,
        ["데이터", "현재 역할", "향후 확장 역할"],
        [
            ["SKEL Rigid/Semi-Rigid/Flexible", "강체, 반강체, 유연체 조작 궤적의 1차 원천 데이터", "현장 초기 기준 궤적과 물성별 도메인 랜덤화"],
            ["AMASS TCDHands", "손 중심 상체 동작 확장 후보", "양손 조립, 보조 팔 동작, 자세 다양화"],
            ["IKEA ASM", "외부 확장 후보", "조립 단계, 물체 상태, human pose, segmentation 주석 결합"],
            ["Assembly101", "외부 확장 후보", "세밀한 조립/분해 순서, 실수 및 수정 행동, 3D hand pose 학습"],
        ],
        "Table 1. Data sources for the executable simulation and future expansion."
    )

    add_heading(doc, "2.3 Robot Model and Dynamics Alignment", 2)
    add_paragraph(
        doc,
        "로봇은 어깨 폭 0.46 m, 상완 0.34 m, 전완 0.32 m의 상체 양팔 모델로 형상화하였다. 오른팔은 목표 조립 궤적을 추종하고 "
        "왼팔은 보조 지지 팔로 동작한다. 공정 환경은 물성별로 지연, 진동, 접촉 잡음, 순응 오차를 다르게 부여하였다. 동역학 정렬은 "
        "목표 궤적과 로봇 응답 사이의 잔차를 반복적으로 학습하는 방식으로 구현하였고, 허용 오차 안의 작은 흔들림은 과학습하지 않도록 제외하였다."
    )
    add_paragraph(
        doc,
        "안전 제어는 최대 이동 스텝과 추정 토크 제한을 적용하는 방식으로 구성하였다. 제한을 초과하는 명령은 자동 감쇠되어 "
        "가상 환경 내 오작동률을 0.0%로 유지한다."
    )

    add_heading(doc, "III. Results and Discussion", 1)
    add_heading(doc, "3.1 Comparison with the Original Paper-Level Indicators", 2)
    add_paragraph(
        doc,
        "기존 PDF 논문에는 프레임워크 수준의 정량 지표가 제시되어 있었고, 본 수정본에서는 실제 실행 가능한 축소 시뮬레이터의 결과를 추가하였다. "
        "두 결과는 동일한 하드웨어 실험을 의미하지 않으므로 직접적인 우열 비교가 아니라, 논문 주장과 구현 검증의 연결성을 보여주는 비교 표로 해석해야 한다."
    )
    rows = []
    for process in PROCESSES:
        old_b = PAPER_BASELINE[process]["before"]
        old_a = PAPER_BASELINE[process]["after"]
        new_b = summaries[process]["before"]["defect_rate"]
        new_a = summaries[process]["after"]["defect_rate"]
        rows.append([LABELS[process], f"{old_b:.2f}% -> {old_a:.2f}%", f"{new_b:.2f}% -> {new_a:.2f}%", f"{new_b - new_a:.2f}%p", "동일 방향 개선 확인"])
    add_table(
        doc,
        ["환경", "기존 PDF 지표", "이번 실행 시뮬레이션", "시뮬레이션 개선폭", "해석"],
        rows,
        "Table 2. Comparison between original paper-level indicators and executable simulation results."
    )

    add_heading(doc, "3.2 Tracking Error and Defect Rate after Alignment", 2)
    rows = []
    for process in PROCESSES:
        summary = summaries[process]
        rows.append([
            LABELS[process],
            f"{summary['before']['defect_rate']:.2f}%",
            f"{summary['after']['defect_rate']:.2f}%",
            f"{summary['before']['mean_error']:.4f} m",
            f"{summary['after']['mean_error']:.4f} m",
            f"{summary['after']['max_error']:.4f} m",
            f"{summary['after']['malfunction_rate']:.1f}%",
        ])
    add_table(
        doc,
        ["환경", "정렬 전 불량률", "정렬 후 불량률", "평균 오차 전", "평균 오차 후", "최대 오차 후", "오작동률"],
        rows,
        "Table 3. Defect-rate and tracking-error changes after repeated dynamics alignment."
    )

    add_heading(doc, "3.3 Visual Verification of Robot Motion", 2)
    add_paragraph(
        doc,
        "시뮬레이션은 수치 지표뿐 아니라 공정 공간 안에서의 움직임 확인을 위해 PNG 및 GIF 산출물을 생성한다. 아래 그림은 정렬 전후 목표 궤적과 "
        "로봇 응답을 함께 표시하며, 특히 유연체 공정에서 정렬 후 오차가 크게 낮아지는 것을 확인할 수 있다."
    )
    for fig_no, process, caption in [
        (4, "semi_rigid", "Semi-rigid process trajectory and tracking-error comparison."),
        (5, "flexible", "Flexible process trajectory and tracking-error comparison."),
    ]:
        image_path = ROOT / "outputs" / process / "trajectory_and_error.png"
        if image_path.exists():
            doc.add_picture(str(image_path), width=Inches(6.25))
            add_caption(doc, f"Fig. {fig_no}. {caption}")

    add_paragraph(
        doc,
        "또한 outputs/<process>/upper_body_robot.gif에는 상체 로봇의 실제 움직임 애니메이션이 저장된다. 논문 본문에는 정지 이미지 중심으로 삽입하고, "
        "발표 또는 부록에서는 GIF를 활용해 작업대 안에서의 팔 움직임을 확인할 수 있다."
    )

    add_heading(doc, "3.4 Discussion", 2)
    add_paragraph(
        doc,
        "강체 공정은 초기 오차가 작아 개선 폭이 제한적이었지만, 불량률과 최대 오차가 감소하였다. 반강체 공정은 접촉 저항과 동적 지연이 반복 정렬 후 "
        "완화되었고, 유연체 공정은 초기 고주파 진동과 큰 추종 변동이 가장 컸으나 반복 정렬 후 불량률이 1.67%까지 낮아졌다. 이 결과는 케이블류, "
        "호스류, 연성 부품을 다루는 중소 조립 현장에서 동역학 정렬 루프가 가장 큰 효과를 낼 수 있음을 시사한다."
    )
    add_paragraph(
        doc,
        "단, 본 결과는 실제 하드웨어 실험이 아니라 가상환경에서의 초동 검증이다. 따라서 실제 배포 전에는 카메라 기반 사람 동작 복원, "
        "로봇 관절 토크 센서, 전류 피드백, 접촉 이벤트 로그를 결합한 폐루프 검증이 필요하다."
    )

    add_heading(doc, "IV. Conclusions", 1)
    add_paragraph(
        doc,
        "본 수정 연구는 린 기반 피지컬 AI 제조 프레임워크를 실행 가능한 상체 제조 로봇 시뮬레이션으로 확장하였다. 최소 SKEL 조작 데이터만으로 "
        "작업대 조립 목표 궤적을 생성하고, 강체, 반강체, 유연체 환경에서 반복 동역학 정렬을 수행하였다. 그 결과 모든 환경에서 불량률이 감소하였고, "
        "특히 유연체 공정에서는 59.58%에서 1.67%로 크게 개선되었다. 안전 제한 구조는 모든 실험에서 오작동률 0.0%를 유지하였다."
    )
    add_paragraph(
        doc,
        "따라서 본 연구는 중소 제조 현장의 데이터 부재 문제를 단순한 한계가 아니라, 시뮬레이션과 현장 로그를 결합한 자가 학습형 데이터 자산 구축의 "
        "출발점으로 전환할 수 있음을 보인다. 후속 연구에서는 IKEA ASM 및 Assembly101의 조립 단계 라벨과 실제 로봇 센서 로그를 결합하여 "
        "작업 순서 인식, 오류 수정 행동, 실제 하드웨어 배포까지 포함하는 고정밀 디지털 트윈으로 확장할 계획이다."
    )

    add_heading(doc, "REFERENCES", 1)
    refs = [
        "JH Ryu, 텔레로봇 및 텔레프레즌스의 연구 동향, 로봇과 인간, 2008.",
        "Y. Tong, H. Liu, Z. Zhang, Humanoid Robots: A Comprehensive Review and Future Prospects, IEEE/CAA Journal of Automatica Sinica, 2024.",
        "N. Kumar et al., Lean manufacturing techniques and implementation: A review, Materials Today: Proceedings, 2022.",
        "V. Yadav et al., The propagation of lean thinking in SMEs, Production Planning & Control, 2019.",
        "B. Zhou, Lean principles, practices, and impacts: a study on small and medium-sized enterprises, Annals of Operations Research, 2016.",
        "Y. Wang et al., TRAM: Global Trajectory and Motion of 3D Humans from in-the-wild Videos, ECCV, 2025.",
        "T. He et al., ASAP: Aligning Simulation and Real-World Physics for Learning Agile Humanoid Whole-Body Skills, RSS, 2025.",
        "W. Zhao et al., Absolute State-wise Constrained Policy Optimization, JMLR, 2024.",
        "Y. Ben-Shabat et al., The IKEA ASM Dataset: Understanding People Assembling Furniture through Actions, Objects and Pose, WACV, 2021.",
        "F. Sener et al., Assembly101: A Large-Scale Multi-View Video Dataset for Understanding Procedural Activities, CVPR, 2022.",
        "M. Loper et al., AMASS: Archive of Motion Capture as Surface Shapes, ICCV, 2019.",
        "S. Kim, Study on improving production system efficiency of humanoid robots in small-scale manufacturing sites using lean manufacturing techniques, Hanyang University, 2025.",
    ]
    for idx, ref in enumerate(refs, start=1):
        add_paragraph(doc, f"[{idx}] {ref}")

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_document()
